"""
core/report.py
Report generation utilities for isolated footing design results.

Supports:
  - Console (print_summary)
  - Excel (openpyxl)
  - Word  (python-docx)
  - PDF   (reportlab)
"""

from __future__ import annotations

from typing import Any, Dict, List, Optional

from core.anchorage import AnchorageResult
from core.optimizer import OptimizationResult
from core.rc_design import DesignResult, MaterialProperties
from core.soil_pressure import FootingGeometry, PressureResult, SoilProperties
from core.stability import StabilityResult


# ---------------------------------------------------------------------------
# Summary dictionary builder
# ---------------------------------------------------------------------------

def generate_summary_dict(
    geom: FootingGeometry,
    soil: SoilProperties,
    materials: MaterialProperties,
    pressure_results: List[PressureResult],
    stability_results: List[StabilityResult],
    design_result: DesignResult,
    anchorage_result: Optional[AnchorageResult] = None,
    optimization_result: Optional[OptimizationResult] = None,
) -> dict:
    """
    Assemble a nested dictionary with every relevant result for reporting.
    """
    summary: Dict[str, Any] = {}

    # --- Geometry -------------------------------------------------------
    summary["geometry"] = {
        "B [m]": geom.B,
        "L [m]": geom.L,
        "h [m]": geom.h,
        "bx [m]": geom.bx,
        "by [m]": geom.by,
        "cover [m]": geom.cover,
        "area [m²]": geom.B * geom.L,
        "volume [m³]": geom.B * geom.L * geom.h,
    }

    # --- Soil -----------------------------------------------------------
    summary["soil"] = {
        "qa [kPa]": soil.qa,
        "γ_soil [kN/m³]": soil.gamma_soil,
        "Df [m]": soil.Df,
    }

    # --- Materials ------------------------------------------------------
    summary["materials"] = {
        "f'c [MPa]": materials.fc,
        "fy [MPa]": materials.fy,
        "φ_flexure": materials.phi_flexure,
        "φ_shear": materials.phi_shear,
    }

    # --- Pressure results -----------------------------------------------
    summary["pressure_results"] = [
        {
            "combo": pr.combo_name,
            "N_total [kN]": round(pr.N_total, 2),
            "q_max [kPa]": round(pr.q_max, 2),
            "q_min [kPa]": round(pr.q_min, 2),
            "ex [m]": round(pr.eccentricity_x, 4),
            "ey [m]": round(pr.eccentricity_y, 4),
            "full_contact": pr.full_contact,
            "contact_ratio": round(pr.contact_ratio, 3),
            "passes_qa": pr.passes_qa,
        }
        for pr in pressure_results
    ]

    # --- Stability results ----------------------------------------------
    summary["stability_results"] = [
        {
            "combo": sr.combo_name,
            "FS_sliding_x": round(sr.FS_sliding_x, 2),
            "FS_sliding_y": round(sr.FS_sliding_y, 2),
            "FS_OT_x": round(sr.FS_overturning_x, 2),
            "FS_OT_y": round(sr.FS_overturning_y, 2),
            "FS_uplift": round(sr.FS_uplift, 2),
            "passes_all": sr.passes_all,
        }
        for sr in stability_results
    ]

    # --- Design result --------------------------------------------------
    dr = design_result
    summary["design"] = {
        "dx [mm]": round(dr.dx, 1),
        "dy [mm]": round(dr.dy, 1),
        "flexure_x": {
            "Mu [kN·m/m]": round(dr.Mu_x, 2),
            "As_req [mm²/m]": round(dr.As_req_x, 1),
            "As_prov [mm²/m]": round(dr.As_prov_x, 1),
            "bar": dr.bar_x,
            "spacing [mm]": round(dr.spacing_x, 1),
            "φMn [kN·m/m]": round(dr.phi_Mn_x, 2),
            "passes": dr.passes_flexure_x,
        },
        "flexure_y": {
            "Mu [kN·m/m]": round(dr.Mu_y, 2),
            "As_req [mm²/m]": round(dr.As_req_y, 1),
            "As_prov [mm²/m]": round(dr.As_prov_y, 1),
            "bar": dr.bar_y,
            "spacing [mm]": round(dr.spacing_y, 1),
            "φMn [kN·m/m]": round(dr.phi_Mn_y, 2),
            "passes": dr.passes_flexure_y,
        },
        "shear_x": {
            "Vu [kN/m]": round(dr.Vu_x, 2),
            "φVc [kN/m]": round(dr.phi_Vc_x, 2),
            "passes": dr.passes_shear_x,
        },
        "shear_y": {
            "Vu [kN/m]": round(dr.Vu_y, 2),
            "φVc [kN/m]": round(dr.phi_Vc_y, 2),
            "passes": dr.passes_shear_y,
        },
        "punching": {
            "Vu2way [kN]": round(dr.Vu2way, 2),
            "φVc2way [kN]": round(dr.phi_Vc2way, 2),
            "passes": dr.passes_punching,
        },
        "passes_all": dr.passes_all,
    }

    # --- Anchorage ------------------------------------------------------
    if anchorage_result is not None:
        ar = anchorage_result
        summary["anchorage"] = {
            "Mu_transfer [kN·m]": round(ar.Mu_transfer, 2),
            "φVn_shear_friction [kN]": round(ar.phi_Vn_transfer, 2),
            "passes_shear_friction": ar.passes_shear_friction,
            "ld_required [mm]": round(ar.ld_required, 1),
            "ld_available [mm]": round(ar.ld_available, 1),
            "passes_development": ar.passes_development,
            "can_be_fixed": ar.can_be_fixed,
            "theta_estimated [mrad]": round(ar.theta_estimated * 1000.0, 3),
            "warnings": ar.warnings,
            "notes": ar.notes,
        }

    # --- Optimization ---------------------------------------------------
    if optimization_result is not None:
        opt = optimization_result
        summary["optimization"] = {
            "converged": opt.converged,
            "reason": opt.reason,
            "n_iterations": opt.n_iterations,
            "n_feasible": opt.n_feasible,
            "objective_value": round(opt.objective_value, 4),
            "best_geometry": {
                "B [m]": opt.best_geometry.B if opt.best_geometry else None,
                "L [m]": opt.best_geometry.L if opt.best_geometry else None,
                "h [m]": opt.best_geometry.h if opt.best_geometry else None,
            } if opt.best_geometry else None,
        }

    return summary


# ---------------------------------------------------------------------------
# Console printer
# ---------------------------------------------------------------------------

def print_summary(summary: dict) -> None:
    """Print formatted summary to console."""

    def _section(title: str) -> None:
        print(f"\n{'='*60}")
        print(f"  {title}")
        print(f"{'='*60}")

    def _kv(key: str, value: Any, indent: int = 2) -> None:
        pad = " " * indent
        ok_str = ""
        if isinstance(value, bool):
            ok_str = " ✓" if value else " ✗"
        print(f"{pad}{key}: {value}{ok_str}")

    _section("FOOTING GEOMETRY")
    for k, v in summary.get("geometry", {}).items():
        _kv(k, v)

    _section("SOIL & MATERIALS")
    for k, v in summary.get("soil", {}).items():
        _kv(k, v)
    for k, v in summary.get("materials", {}).items():
        _kv(k, v)

    _section("CONTACT PRESSURE RESULTS")
    for pr in summary.get("pressure_results", []):
        print(f"  [{pr['combo']}]  q_max={pr['q_max [kPa]']} kPa  "
              f"q_min={pr['q_min [kPa]']} kPa  "
              f"passes={'✓' if pr['passes_qa'] else '✗'}")

    _section("STABILITY RESULTS")
    for sr in summary.get("stability_results", []):
        print(f"  [{sr['combo']}]  "
              f"FS_sl_x={sr['FS_sliding_x']}  "
              f"FS_sl_y={sr['FS_sliding_y']}  "
              f"FS_OT_x={sr['FS_OT_x']}  "
              f"FS_OT_y={sr['FS_OT_y']}  "
              f"passes={'✓' if sr['passes_all'] else '✗'}")

    _section("STRUCTURAL DESIGN")
    dr = summary.get("design", {})
    print(f"  Effective depths: dx={dr.get('dx [mm]')} mm, dy={dr.get('dy [mm]')} mm")
    for direction in ("flexure_x", "flexure_y"):
        d = dr.get(direction, {})
        print(f"  {direction}: Mu={d.get('Mu [kN·m/m]')} kN·m/m  "
              f"bar={d.get('bar')} @{d.get('spacing [mm]')} mm  "
              f"φMn={d.get('φMn [kN·m/m]')} kN·m/m  "
              f"passes={'✓' if d.get('passes') else '✗'}")
    for shear in ("shear_x", "shear_y"):
        d = dr.get(shear, {})
        print(f"  {shear}: Vu={d.get('Vu [kN/m]')} kN/m  "
              f"φVc={d.get('φVc [kN/m]')} kN/m  "
              f"passes={'✓' if d.get('passes') else '✗'}")
    pun = dr.get("punching", {})
    print(f"  punching: Vu={pun.get('Vu2way [kN]')} kN  "
          f"φVc={pun.get('φVc2way [kN]')} kN  "
          f"passes={'✓' if pun.get('passes') else '✗'}")
    print(f"  OVERALL DESIGN: {'PASS ✓' if dr.get('passes_all') else 'FAIL ✗'}")

    if "anchorage" in summary:
        _section("ANCHORAGE / MOMENT TRANSFER")
        for k, v in summary["anchorage"].items():
            if k != "warnings":
                _kv(k, v)
        for w in summary["anchorage"].get("warnings", []):
            print(f"  ⚠ {w}")

    if "optimization" in summary:
        _section("OPTIMIZATION")
        opt = summary["optimization"]
        print(f"  Converged: {opt['converged']}  ({opt['reason']})")
        print(f"  Evaluated {opt['n_iterations']} combos, {opt['n_feasible']} feasible")
        bg = opt.get("best_geometry")
        if bg:
            print(f"  Best: B={bg['B [m]']} m, L={bg['L [m]']} m, h={bg['h [m]']} m")
            print(f"  Objective value: {opt['objective_value']}")


# ---------------------------------------------------------------------------
# Excel export
# ---------------------------------------------------------------------------

def export_to_excel(summary: dict, filepath: str) -> None:
    """Export the complete summary to an Excel workbook."""
    try:
        import openpyxl
        from openpyxl.styles import Font, PatternFill, Alignment
    except ImportError:
        raise ImportError("openpyxl is required for Excel export.")

    wb = openpyxl.Workbook()

    # ---- Helper --------------------------------------------------------
    def _write_sheet(ws, data_dict: dict, title: str) -> None:
        ws.title = title
        header_font = Font(bold=True, color="FFFFFF")
        header_fill = PatternFill("solid", fgColor="1F4E79")
        row = 1
        for key, val in _flatten(data_dict).items():
            ws.cell(row=row, column=1, value=key).font = Font(bold=True)
            ws.cell(row=row, column=2, value=str(val))
            row += 1
        ws.column_dimensions["A"].width = 35
        ws.column_dimensions["B"].width = 25

    def _flatten(d: dict, parent_key: str = "", sep: str = " / ") -> dict:
        items: list = []
        for k, v in d.items():
            new_key = f"{parent_key}{sep}{k}" if parent_key else k
            if isinstance(v, dict):
                items.extend(_flatten(v, new_key, sep).items())
            else:
                items.append((new_key, v))
        return dict(items)

    # ---- Sheets --------------------------------------------------------
    ws_geom = wb.active
    _write_sheet(ws_geom, summary.get("geometry", {}), "Geometry")

    for section, title in [
        ("soil", "Soil"),
        ("materials", "Materials"),
        ("design", "Design"),
    ]:
        ws = wb.create_sheet(title)
        _write_sheet(ws, summary.get(section, {}), title)

    # Pressure table
    ws_pr = wb.create_sheet("Pressure")
    prs = summary.get("pressure_results", [])
    if prs:
        headers = list(prs[0].keys())
        for col, h in enumerate(headers, 1):
            cell = ws_pr.cell(row=1, column=col, value=h)
            cell.font = Font(bold=True)
        for row, pr in enumerate(prs, 2):
            for col, key in enumerate(headers, 1):
                ws_pr.cell(row=row, column=col, value=pr[key])

    # Stability table
    ws_st = wb.create_sheet("Stability")
    srs = summary.get("stability_results", [])
    if srs:
        headers = list(srs[0].keys())
        for col, h in enumerate(headers, 1):
            ws_st.cell(row=1, column=col, value=h).font = Font(bold=True)
        for row, sr in enumerate(srs, 2):
            for col, key in enumerate(headers, 1):
                ws_st.cell(row=row, column=col, value=sr[key])

    wb.save(filepath)
    print(f"Excel report saved: {filepath}")


# ---------------------------------------------------------------------------
# Word export
# ---------------------------------------------------------------------------

def export_to_docx(summary: dict, filepath: str) -> None:
    """Export the summary to a Word document."""
    try:
        from docx import Document
        from docx.shared import Pt, RGBColor
    except ImportError:
        raise ImportError("python-docx is required for Word export.")

    doc = Document()
    doc.add_heading("Isolated Footing Design Report", level=0)

    def _add_table(doc, data: dict, title: str) -> None:
        doc.add_heading(title, level=1)
        table = doc.add_table(rows=1, cols=2)
        table.style = "Light Grid Accent 1"
        hdr = table.rows[0].cells
        hdr[0].text = "Parameter"
        hdr[1].text = "Value"
        for k, v in data.items():
            row = table.add_row().cells
            row[0].text = str(k)
            row[1].text = str(v)

    _add_table(doc, summary.get("geometry", {}), "Geometry")
    _add_table(doc, summary.get("soil", {}), "Soil Properties")
    _add_table(doc, summary.get("materials", {}), "Material Properties")
    _add_table(doc, summary.get("design", {}).get("flexure_x", {}), "Flexure – X Direction")
    _add_table(doc, summary.get("design", {}).get("flexure_y", {}), "Flexure – Y Direction")
    _add_table(doc, summary.get("design", {}).get("shear_x", {}), "One-Way Shear – X")
    _add_table(doc, summary.get("design", {}).get("shear_y", {}), "One-Way Shear – Y")
    _add_table(doc, summary.get("design", {}).get("punching", {}), "Punching Shear")

    doc.save(filepath)
    print(f"Word report saved: {filepath}")


# ---------------------------------------------------------------------------
# PDF export
# ---------------------------------------------------------------------------

def export_to_pdf(summary: dict, filepath: str) -> None:
    """Export the summary to PDF using reportlab."""
    try:
        from reportlab.lib import colors
        from reportlab.lib.pagesizes import A4
        from reportlab.lib.styles import getSampleStyleSheet
        from reportlab.lib.units import cm
        from reportlab.platypus import (
            Paragraph,
            SimpleDocTemplate,
            Spacer,
            Table,
            TableStyle,
        )
    except ImportError:
        raise ImportError("reportlab is required for PDF export.")

    doc_obj = SimpleDocTemplate(filepath, pagesize=A4)
    styles = getSampleStyleSheet()
    story = []

    story.append(Paragraph("Isolated Footing Design Report", styles["Title"]))
    story.append(Spacer(1, 0.5 * cm))

    def _section_table(title: str, data: dict) -> None:
        story.append(Paragraph(title, styles["Heading2"]))
        rows = [["Parameter", "Value"]] + [[str(k), str(v)] for k, v in data.items()]
        t = Table(rows, colWidths=[9 * cm, 8 * cm])
        t.setStyle(
            TableStyle([
                ("BACKGROUND", (0, 0), (-1, 0), colors.HexColor("#1F4E79")),
                ("TEXTCOLOR",  (0, 0), (-1, 0), colors.white),
                ("FONTNAME",   (0, 0), (-1, 0), "Helvetica-Bold"),
                ("ROWBACKGROUNDS", (0, 1), (-1, -1), [colors.white, colors.HexColor("#EBF3FB")]),
                ("GRID",       (0, 0), (-1, -1), 0.5, colors.grey),
                ("FONTSIZE",   (0, 0), (-1, -1), 8),
            ])
        )
        story.append(t)
        story.append(Spacer(1, 0.3 * cm))

    _section_table("Geometry", summary.get("geometry", {}))
    _section_table("Soil Properties", summary.get("soil", {}))
    _section_table("Material Properties", summary.get("materials", {}))

    dr = summary.get("design", {})
    for key in ("flexure_x", "flexure_y", "shear_x", "shear_y", "punching"):
        if key in dr:
            _section_table(f"Design – {key.replace('_', ' ').title()}", dr[key])

    doc_obj.build(story)
    print(f"PDF report saved: {filepath}")
